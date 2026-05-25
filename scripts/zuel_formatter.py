# -*- coding: utf-8 -*-
"""
Markdown to Academic PDF conversion script.
Created for the markdown-to-academic-pdf agent skill.
"""
import os
import sys
import argparse
import subprocess
import io
import json
# Imports are deferred to individual functions to prevent missing dependency errors when not used.

# Hardcoded standard Microsoft Edge executable path on Windows
EDGE_PATH = "C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe"

ACADEMIC_CSS_TEMPLATE = """
/* A4 Page Layout & Margins */
@page {{
    size: A4;
    margin-top: {margin_top}cm;
    margin-bottom: {margin_bottom}cm;
    margin-left: {margin_left}cm;
    margin-right: {margin_right}cm;
}}

body {{
    font-family: 'Times New Roman', 'SimSun', 'STSong', serif;
    color: #000000;
    line-height: 20pt; /* 固定值 20 磅 */
    margin: 0;
    padding: 0;
    font-size: 12pt; /* 小四号 */
}}

/* 正文段落：首行缩进 2 字符，两端对齐，段前后 0 行 */
p {{
    text-indent: 2em; /* 首行缩进 2 字符 */
    margin-top: 0;
    margin-bottom: 0;
    text-align: justify;
    line-height: 20pt;
}}

p strong {{
    font-weight: bold;
}}

/* 列表与引用重置缩进 */
ul, ol {{
    margin-top: 0;
    margin-bottom: 0;
    padding-left: 2em;
    text-indent: 0 !important;
}}
li {{
    text-indent: 0 !important;
    line-height: 20pt; /* 固定值 20 磅 */
}}
blockquote p {{
    text-indent: 0 !important;
}}

/* 一级标题 (顶格，黑体，三号, 段前 24 磅，段后 18 磅) */
h2 {{
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 16pt; /* 三号 */
    text-align: left;
    font-weight: bold;
    margin-top: 24pt; /* 段前 24 磅 */
    margin-bottom: 18pt; /* 段后 18 磅 */
    text-indent: 0 !important;
    page-break-after: avoid;
}}

/* 二级标题 (顶格，宋体加粗，四号, 段前 18 磅，段后 12 磅) */
h3 {{
    font-family: 'Times New Roman', 'SimSun', 'STSong', serif;
    font-size: 14pt; /* 四号 */
    text-align: left;
    font-weight: bold;
    margin-top: 18pt; /* 段前 18 磅 */
    margin-bottom: 12pt; /* 段后 12 磅 */
    text-indent: 0 !important;
    page-break-after: avoid;
}}

/* 三级标题 (顶格，宋体加粗，小四号) */
h4 {{
    font-family: 'Times New Roman', 'SimSun', 'STSong', serif;
    font-size: 12pt; /* 小四号 */
    text-align: left;
    font-weight: bold;
    margin-top: 12pt;
    margin-bottom: 6pt;
    text-indent: 0 !important;
    page-break-after: avoid;
}}

/* 封面排版 CSS (符合 ZUEL 48pt 校名与 20pt 个人信息及校徽要求) */
.cover-page {{
    page-break-after: always;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: flex-start;
    height: 23.5cm; /* 限制高度防止溢出 */
    padding-top: 1.5cm;
    box-sizing: border-box;
    text-align: center;
}}
.cover-school {{
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 48pt; /* 中南财经政法大学字号大号 */
    font-weight: bold;
    letter-spacing: 0.1em;
    margin-top: 0.5cm;
    margin-bottom: 0.3cm;
}}
.cover-paper-type {{
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 48pt; /* 本科生结课报告字号大号 */
    font-weight: bold;
    letter-spacing: 0.1em;
    margin-bottom: 1.0cm;
}}
.cover-logo {{
    margin-top: 0.2cm;
    margin-bottom: 1.2cm;
    display: flex;
    justify-content: center;
    align-items: center;
}}
.cover-logo img {{
    width: 5.5cm;
    height: 5.5cm;
    display: block;
}}
.cover-info-table {{
    width: 70%;
    margin: 0.5cm auto 0 auto;
}}
.cover-info-table table {{
    width: 100%;
    border-collapse: collapse;
    border: none !important;
    margin: 0;
}}
.cover-info-table td {{
    border: none !important;
    padding: 10px 0;
    font-family: 'Kaiti', 'STKaiti', serif;
    font-size: 20pt; /* 信息字号二号/小二 */
    text-align: left;
    line-height: 1.4;
}}
.cover-info-table td.info-label {{
    width: 38%;
    text-align: left;
    font-weight: bold;
    white-space: nowrap;
}}
.cover-info-table td.info-val {{
    width: 62%;
    border-bottom: 2.0px solid #000000 !important;
    text-align: left;
    padding-left: 10px;
}}

.main-paper-title {{
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 16pt; /* 三号黑体加粗 */
    text-align: center;
    font-weight: bold;
    margin-top: 24pt;
    margin-bottom: 18pt;
    text-indent: 0 !important;
    page-break-before: always; /* 独立成页的起点 */
}}

/* 摘要排版 CSS (符合 ZUEL 16pt 摘要/关键词要求) */
.abstract-page {{
    page-break-after: always;
    padding-top: 1.5cm;
    box-sizing: border-box;
}}
.abstract-page h2 {{
    text-align: center !important;
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 16pt; /* 摘要标题三号 (同步一级标题) */
    font-weight: bold;
    margin-bottom: 1.2cm;
    text-indent: 0 !important;
}}
.abstract-page p {{
    margin-bottom: 1.5cm;
}}
.abstract-page .keywords-line {{
    font-size: 12pt;
    line-height: 20pt; /* 固定值 20 磅 */
    text-indent: 0 !important;
    margin-top: 1cm;
}}
.abstract-page .keywords-label {{
    font-size: 16pt; /* 关键词标签三号 (同步一级标题) */
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-weight: bold;
}}

/* 目录排版 CSS (符合 ZUEL 16pt 目录要求) */
.toc-page {{
    page-break-after: always;
    padding-top: 1.5cm;
    box-sizing: border-box;
}}
.toc-page h2 {{
    text-align: center !important;
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 16pt; /* 目录标题三号 */
    font-weight: bold;
    margin-bottom: 1.5cm;
    text-indent: 0 !important;
}}
.toc-page ul {{
    list-style-type: none;
    padding-left: 0;
    line-height: 2.2;
}}
.toc-page li {{
    font-family: 'SimSun', 'STSong', serif;
    font-size: 12pt;
    margin-bottom: 8px;
    border-bottom: 1px dashed #cccccc;
    display: flex;
    justify-content: space-between;
    text-indent: 0 !important;
}}
.toc-page li.level-1 {{
    font-size: 14pt; /* 目录一级标题四号 */
    font-weight: bold;
}}
.toc-page li.level-2 {{
    font-size: 12pt; /* 目录二级标题小四 */
    padding-left: 1.5em;
}}
.toc-page a {{
    text-decoration: none;
    color: #000000;
}}

/* 参考文献 (悬挂缩进) */
.references {{
    margin-top: 10px;
}}
.references-title {{
    font-family: 'SimHei', 'STHeiti', sans-serif;
    font-size: 16pt; /* 参考文献标题三号 */
    text-align: center;
    font-weight: bold;
    margin-top: 24pt;
    margin-bottom: 18pt;
    text-indent: 0 !important;
}}
.references p {{
    font-family: 'Times New Roman', 'SimSun', 'STSong', serif;
    font-size: 12pt; /* 小四号 */
    line-height: 20pt; /* 固定值 20 磅 */
    text-indent: -2em !important;
    padding-left: 2em;
    margin-bottom: 8px;
    text-align: justify;
}}

/* Table spacing */
table {{
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0 25px 0;
    font-size: 10.5pt; /* 表格五号 */
    text-indent: 0 !important;
    line-height: 1.4;
}}
th, td {{
    border: 1px solid #000000;
    padding: 8px 10px;
    text-indent: 0 !important;
}}
th {{
    background-color: #f2f2f2;
    font-weight: bold;
}}

/* Image formatting */
img {{
    max-width: 90%;
    height: auto;
    display: block;
    margin: 15px auto 10px auto;
    text-indent: 0 !important;
}}

/* Code blocks */
pre {{
    background-color: #fafafa;
    border: 0.5px solid #bbbbbb;
    padding: 10px 15px;
    border-radius: 4px;
    overflow-x: auto;
    font-family: Consolas, Monaco, 'Courier New', monospace;
    font-size: 10pt;
    margin: 15px 0;
    line-height: 14pt;
    text-indent: 0 !important;
}}
code {{
    font-family: Consolas, Monaco, 'Courier New', monospace;
    font-size: 10.5pt;
}}
pre code {{
    font-size: 10pt;
}}

@media print {{
    body {{
        padding: 0;
    }}
    h1, h2, h3, h4 {{
        page-break-after: avoid;
    }}
    pre, table, img {{
        page-break-inside: avoid;
    }}
}}
"""

def get_logo_base64():
    """获取校徽图片的 Base64 编码"""
    import base64
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "..", "resources", "zuel_logo.png")
    if os.path.exists(logo_path):
        try:
            with open(logo_path, "rb") as f:
                return base64.b64encode(f.read()).decode("utf-8")
        except Exception as e:
            print(f"Error reading logo for base64: {e}", file=sys.stderr)
    return ""

def generate_html(md_path, margin_top=3.0, margin_bottom=2.5, margin_left=3.0, margin_right=2.5):
    """读入MD，转换为Styled HTML"""
    import markdown
    import re
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
        
    # 解析元数据
    metadata = parse_frontmatter(md_path)
    
    # 剥离 YAML frontmatter 以防止作为正文渲染出来
    if md_text.startswith("---"):
        parts = md_text.split("---", 2)
        if len(parts) >= 3:
            md_text = parts[2]
            
    # 清理图表题注，去除英文点号分隔符等，规范为 "图1 XXX" / "表1 XXX"
    md_text = re.sub(r'^(图|表)\s*(\d+)\s*[\.、。．\.]?\s*(.*)$', r'\1\2 \3', md_text, flags=re.MULTILINE)
            
    # md_in_html 扩展必不可少，用于解析封面和承诺书 div 内的 Markdown
    html_body = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'md_in_html'])
    
    # 构建 HTML 封面页并前置注入（预设空数据以满足要求）
    logo_base64 = get_logo_base64()
    logo_src = f"data:image/png;base64,{logo_base64}" if logo_base64 else "zuel_logo.png"
    
    cover_html = f"""
    <div class="cover-page">
        <div class="cover-school">中南财经政法大学</div>
        <div class="cover-paper-type">本科生结课报告</div>
        <div class="cover-logo">
            <img src="{logo_src}">
        </div>
        <div class="cover-info-table">
            <table>
                <tr>
                    <td class="info-label">课程编号：</td>
                    <td class="info-val">&nbsp;</td>
                </tr>
                <tr>
                    <td class="info-label">课程名称：</td>
                    <td class="info-val">&nbsp;</td>
                </tr>
                <tr>
                    <td class="info-label">授课教师：</td>
                    <td class="info-val">&nbsp;</td>
                </tr>
                <tr>
                    <td class="info-label">学生学号：</td>
                    <td class="info-val">&nbsp;</td>
                </tr>
                <tr>
                    <td class="info-label">学生姓名：</td>
                    <td class="info-val">&nbsp;</td>
                </tr>
            </table>
        </div>
    </div>
    """
    
    html_body = cover_html + html_body
    
    # 在 HTML 正文 the 第一个一级标题前插入大标题
    h1_title = f'<h1 class="main-paper-title">{metadata["title"]}</h1>'
    html_body, _ = re.subn(r'(<h2>一、)', f'{h1_title}\\1', html_body, count=1)
    
    # 动态包裹参考文献段落为 class="references" 的 div
    html_body = html_body.replace('<h2>7. 参考文献</h2>', '<h2 class="references-title">7. 参考文献</h2><div class="references">')
    # 处理参考文献与附录的分割
    html_body = html_body.replace('<h2>8. 附录：', '</div><h2>8. 附录：')
    html_body = html_body.replace('<h2>8. 附录</h2>', '</div><h2>8. 附录</h2>')
    
    css_content = ACADEMIC_CSS_TEMPLATE.format(
        margin_top=margin_top,
        margin_bottom=margin_bottom,
        margin_left=margin_left,
        margin_right=margin_right
    )
    
    styled_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Converted Document</title>
<style>
{css_content}
</style>
</head>
<body>
    {html_body}
    
    <script>
    // 自动寻找图表题注并格式化样式（居中、无首行缩进、五号10.5pt加粗）
    document.querySelectorAll('p').forEach(p => {{
        const text = p.textContent.trim();
        if (/^(图|表)\\d+\\s/.test(text)) {{
            p.style.textIndent = '0';
            p.style.textAlign = 'center';
            p.style.fontWeight = 'bold';
            p.style.fontSize = '10.5pt';
            p.style.marginTop = '6px';
            p.style.marginBottom = '12px';
        }}
    }});
    
    // 自动给目录、结语、主要参考文献设置分页，中文双字独立页面大标题保留空格，不含Abstract
    document.querySelectorAll('h2').forEach(h2 => {{
        const text = h2.textContent.trim().replace(/\\s+/g, '');
        if (text === "目录" || text === "结语" || text === "主要参考文献") {{
            h2.style.pageBreakBefore = 'always';
        }}
    }});
    </script>
</body>
</html>
"""
    return styled_html

def post_process_pdf(pdf_path, header_text=None, margin_left=3.0, margin_right=2.5):
    """后处理 PDF 文件，添加学术页脚和页眉页码"""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        from PyPDF2 import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if not os.path.exists(pdf_path):
        print(f"Error: Output PDF file {pdf_path} was not found for post-processing.", file=sys.stderr)
        return

    # 1. 寻找和注册 SimSun (宋体) 字体
    font_path = "C:/Windows/Fonts/simsun.ttc"
    font_name = "Helvetica"
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("SimSun", font_path))
            font_name = "SimSun"
        except Exception as e:
            print(f"ReportLab failed to register SimSun: {e}", file=sys.stderr)

    # 2. 读取 PDF 结构
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    num_pages = len(reader.pages)

    # 3. 动态寻找正文开始的页面 (去除空白字符以实现鲁棒的文本匹配)
    body_start_idx = 1
    for idx in range(num_pages):
        page = reader.pages[idx]
        text = page.extract_text() or ""
        clean_text = "".join(text.split())
        if "1.绪论" in clean_text or "1.绪论与研究背景" in clean_text:
            body_start_idx = idx
            break

    # 4. A4 纸点尺寸计算
    width, height = A4
    margin_left_pts = margin_left * 28.346
    margin_right_pts = margin_right * 28.346

    # 逐页处理，添加水印/印章（页脚和页眉）
    for idx in range(num_pages):
        page = reader.pages[idx]
        
        # 封面、承诺书等前置页面不添加页眉和页脚页码
        if idx < body_start_idx:
            writer.add_page(page)
            continue
            
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=A4)
        
        # 如果指定了 --header 文本，则绘制页眉
        if header_text:
            can.setFont(font_name, 9)
            can.drawCentredString(width / 2.0, height - 51.0, header_text)
            
            can.setLineWidth(0.5)
            can.line(margin_left_pts, height - 56.0, width - margin_right_pts, height - 56.0)
            
        # 绘制页脚页码 (小五号宋体 9pt，居中，距底端 1.5cm = 42.5pt)
        page_num = idx - body_start_idx + 1
        can.setFont(font_name, 9)
        can.drawCentredString(width / 2.0, 42.5, str(page_num))
        
        can.save()
        
        # 合并当前页与绘制的页眉页脚
        packet.seek(0)
        watermark_pdf = PdfReader(packet)
        watermark_page = watermark_pdf.pages[0]
        
        page.merge_page(watermark_page)
        writer.add_page(page)
        
    # 保存后处理的文件
    temp_pdf_path = pdf_path + ".processed"
    with open(temp_pdf_path, "wb") as f:
        writer.write(f)
        
    # 用新文件替换原文件
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    os.rename(temp_pdf_path, pdf_path)
    print("Academic PDF header/footer successfully post-processed.")

def parse_frontmatter(md_path):
    """从 Markdown 头部解析 YAML 元数据"""
    metadata = {
        "title": "（论文题目）",
        "course_name": "（课程名称）",
        "course_code": "（课程编号）",
        "teacher": "（授课教师）",
        "student_id": "                  ",
        "student_name": "                  "
    }
    if not os.path.exists(md_path):
        return metadata
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            yaml_text = parts[1]
            for line in yaml_text.splitlines():
                if ":" in line:
                    key, val = line.split(":", 1)
                    key = key.strip().lower()
                    val = val.strip().strip('"').strip("'")
                    if key in ["title", "题目"]:
                        metadata["title"] = val
                    elif key in ["course_name", "课程名称"]:
                        metadata["course_name"] = val
                    elif key in ["course_code", "课程编号"]:
                        metadata["course_code"] = val
                    elif key in ["teacher", "授课教师"]:
                        metadata["teacher"] = val
                    elif key in ["student_id", "学生学号"]:
                        metadata["student_id"] = val
                    elif key in ["student_name", "学生姓名"]:
                        metadata["student_name"] = val
    return metadata

def generate_docx_from_markdown(md_path, output_docx_path):
    """将 Markdown 编译为完全符合 ZUEL 学术格式规范的 DOCX 文档"""
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx.enum.table
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    # 1. 运行 Pandoc 将 Markdown 编译为基础 DOCX，使用 ZUEL 模版作为样式参考
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, "..", "resources", "zuel_template.docx")
    cmd = ["pandoc", "-s", md_path, "-o", output_docx_path]
    if os.path.exists(template_path):
        cmd.append(f"--reference-doc={template_path}")
    print("Running pandoc to convert Markdown to base DOCX...")
    subprocess.run(cmd, check=True)
    
    # 2. 解析元数据
    metadata = parse_frontmatter(md_path)
    
    # 3. 使用 python-docx 加载文档并进行后处理（添加封面、摘要级次以及微调正文段落）
    doc = docx.Document(output_docx_path)
    
    # 设置标准的 ZUEL 页边距：上3.0cm 下2.5cm 左3.0cm 右2.5cm
    for section in doc.sections:
        section.top_margin = Inches(3.0 / 2.54)
        section.bottom_margin = Inches(2.5 / 2.54)
        section.left_margin = Inches(3.0 / 2.54)
        section.right_margin = Inches(2.5 / 2.54)
        
    first_p = doc.paragraphs[0] if doc.paragraphs else None
    if not first_p:
        first_p = doc.add_paragraph()
        
    # 定义设置中西文字体的辅助函数
    def set_run_font(run, font_name, size_pt=None, bold=None):
        if size_pt:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        run.font.name = font_name
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)

    # 4. 在文档最前面，按顺序插入封面页各元素
    # 我们以 first_p 为锚点，通过 insert_paragraph_before 按从上到下的顺序插入：
    # 1) p_school
    p_school = first_p.insert_paragraph_before("")
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school.paragraph_format.space_before = Pt(40)  # 上边距
    p_school.paragraph_format.space_after = Pt(12)
    r_school = p_school.add_run("中南财经政法大学")
    set_run_font(r_school, "黑体", 48, True)
    
    # 2) p_title
    p_title = first_p.insert_paragraph_before("")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(36)
    r_title = p_title.add_run("本科生结课报告")
    set_run_font(r_title, "黑体", 48, True)
    
    # 3) p_logo (插入校徽图片)
    p_logo = first_p.insert_paragraph_before("")
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(48)
    r_logo = p_logo.add_run()
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, "..", "resources", "zuel_logo.png")
    if os.path.exists(logo_path):
        r_logo.add_picture(logo_path, width=Inches(2.2))
    else:
        print(f"Warning: Logo image not found at {logo_path}", file=sys.stderr)
        
    # 4) p_space_table (空白段落以拉开表格距离)
    p_space_table = first_p.insert_paragraph_before("")
    p_space_table.paragraph_format.space_before = Pt(24)
    p_space_table.paragraph_format.space_after = Pt(12)
    
    # 5) table (课程信息表格 5行2列，左对齐，带底边框下划线)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Normal Table'
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    first_p._p.addprevious(table._tbl)
    
    labels = ["课程编号：", "课程名称：", "授课教师：", "学生学号：", "学生姓名："]
    vals = [
        "",
        "",
        "",
        "",
        ""
    ]
    
    def set_cell_border_bottom(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5 pt
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        
    for r_idx in range(5):
        row = table.rows[r_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        # 标签单元格 (左对齐)
        cell_lbl.text = ""
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_lbl = p_lbl.add_run(labels[r_idx])
        set_run_font(r_lbl, "楷体", 20, True)
        
        # 数值单元格 (左对齐)
        cell_val.text = ""
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_val = p_val.add_run(vals[r_idx])
        set_run_font(r_val, "楷体", 20, False)
        
        # 设置数值单元格下边框为下划线
        set_cell_border_bottom(cell_val)
        
        # 设定单元格宽度
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(3.5)
        
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(3.5)
        
    # 6) p_break (分页符)
    p_break = first_p.insert_paragraph_before("")
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after = Pt(0)
    p_break.paragraph_format.page_break_before = True

    # 查找并移动论文的大标题 (由 Pandoc 在 first_p 处生成)
    title_text = metadata.get("title", "").strip()
    is_title_p = (first_p and first_p.text.strip() == title_text)
    
    if is_title_p and title_text:
        # 1. 从原位置删除 title 段落
        p_element = first_p._p
        p_element.getparent().remove(p_element)
        
        # 2. 查找第一个正文一级标题 (如 "一、引言")，使用状态机排除目录项
        first_h1_p = None
        in_toc_section = False
        seen_first_level_headings = set()
        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not p_text:
                continue
            clean_p = "".join(p_text.split())
            if clean_p == "目录":
                in_toc_section = True
                continue
            if in_toc_section:
                if re.match(r"^[一二三四五六七八九十]+、", p_text):
                    if p_text in seen_first_level_headings:
                        in_toc_section = False
                        first_h1_p = p
                        break
                    else:
                        seen_first_level_headings.add(p_text)
            else:
                if re.match(r"^[一二三四五六七八九十]+、", p_text):
                    first_h1_p = p
                    break
                
        # 3. 在该一级标题前插入大标题段落
        if first_h1_p:
            p_new_title = first_h1_p.insert_paragraph_before("")
            p_new_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_new_title.paragraph_format.space_before = Pt(24)
            p_new_title.paragraph_format.space_after = Pt(18)
            p_new_title.paragraph_format.page_break_before = True
            p_new_title.paragraph_format.first_line_indent = Pt(0)
            r_new_title = p_new_title.add_run(title_text)
            set_run_font(r_new_title, "黑体", 16, True) # 16pt (三号) 黑体加粗

    # 5. 遍历文档正文的所有段落，强制套用规范样式
    print("Formatting paragraphs and titles inside DOCX to match ZUEL specifications...")
    in_toc_section = False
    seen_first_level_headings = set()

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # 跳过我们刚刚在封面插入的包含特有字号的元素和大标题
        if text == title_text or any(r.font.size in [Pt(26), Pt(48), Pt(20)] for r in p.runs):
            continue
            
        clean_text = "".join(text.split())
        is_目录_title = (clean_text == "目录")

        # 如果已经在目录区域中，且遇到了第一个一级标题，判断是否为正文开始
        if in_toc_section and not is_目录_title:
            if re.match(r"^[一二三四五六七八九十]+、", text):
                if text in seen_first_level_headings:
                    in_toc_section = False
                else:
                    seen_first_level_headings.add(text)
            
        # 识别并处理目录项，避免被误判为标题，且目录项不应首行缩进
        is_toc_entry = (in_toc_section and not is_目录_title) or "TOC" in p.style.name or p.style.name.startswith("toc") or p.style.name.startswith("TOC")
        if not is_toc_entry:
            if (p.style.name == "Normal" or "Hyperlink" in p.style.name or "List" in p.style.name) and (re.search(r"\d+$", text) or "..." in text or "\t" in text or "•" in text):
                is_toc_entry = True
                
        if is_toc_entry:
            p.paragraph_format.line_spacing = Pt(20)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs:
                set_run_font(r, "宋体", 12, False)
            continue

        if is_目录_title:
            in_toc_section = True
            
        # A. 一级标题 "一、" (16pt 黑体加粗)
        if re.match(r"^[一二三四五六七八九十]+、", text):
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            # 一级标题不再默认分页；主要大标题均已在其单独的分支处理了分页
            p.paragraph_format.page_break_before = False
            for r in p.runs:
                set_run_font(r, "黑体", 16, True)
                
        # B. 二级标题 "（一）" (14pt 宋体加粗)
        elif re.match(r"^（[一二三四五六七八九十]+）", text):
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            for r in p.runs:
                set_run_font(r, "宋体", 14, True)
                
        # C. 三级标题 "1." (12pt 宋体加粗)
        elif re.match(r"^\d+[\.、]", text):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                set_run_font(r, "宋体", 12, True)
                
        # D. 摘要/目录/结语/参考文献 大标题 (16pt 黑体/Times New Roman 加粗，居中对齐，单独页，保留空格)
        elif clean_text in ["摘要", "Abstract", "目录", "主要参考文献", "结语"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            # 英文摘要 Abstract 紧接着中文摘要，不进行分页；其它大标题（中文摘要、目录、结语、主要参考文献）强制分页
            p.paragraph_format.page_break_before = (clean_text != "Abstract")
            p.paragraph_format.first_line_indent = Pt(0)  # 大标题不缩进
            for r in p.runs:
                set_run_font(r, "黑体" if clean_text != "Abstract" else "Times New Roman", 16, True)
                
        # E. 关键词段落 (中文/英文，标签 16pt 黑体加粗，内容 12pt 宋体)
        elif "关键词" in text or "Keywords" in text or "Key words" in text or "Key Words" in text:
            p.paragraph_format.line_spacing = Pt(20)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Pt(0)  # 关键词行不缩进
            for r in p.runs:
                r_text = r.text.strip()
                if "关键词" in r_text or "Keywords" in r_text or "Key words" in r_text or "Key Words" in r_text:
                    set_run_font(r, "黑体" if "关键词" in r_text else "Times New Roman", 16, True)
                else:
                    set_run_font(r, "宋体" if "关键词" in text else "Times New Roman", 12, False)

        # G. 图表题注格式化 (表1 XXX, 图1 XXX - 去点、居中、五号宋体加粗、不缩进)
        elif re.match(r'^(表|图)\s*\d+', text):
            m_cap = re.match(r'^(表|图)\s*(\d+)([\.、。．\.]*\s*)(.*)', text)
            if m_cap:
                label = m_cap.group(1)
                num = m_cap.group(2)
                rest = m_cap.group(4).strip()
                
                p.text = f"{label}{num} {rest}"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(20) # 20磅行距
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Pt(0) # 题注不缩进
                
                for r in p.runs:
                    set_run_font(r, "宋体", 10.5, True) # 10.5pt (五号) 宋体加粗

        # F. 正文普通段落
        else:
            p.paragraph_format.line_spacing = Pt(20) # 固定值 20 磅
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # 首行缩进 2 字符 (约 24pt)
            if not (text.startswith("  ") or text.startswith("\u3000\u3000")):
                p.paragraph_format.first_line_indent = Pt(24)
                
            for r in p.runs:
                set_run_font(r, "宋体", 12, False) # 12pt (小四) 宋体/新罗马
                
    doc.save(output_docx_path)

def convert_cmd(args):
    """convert 子命令处理函数，支持双格式同时输出"""
    md_path = os.path.abspath(args.input)
    if not os.path.exists(md_path):
        print(f"Error: Input Markdown file does not exist: {md_path}", file=sys.stderr)
        sys.exit(1)

    # 确定输出格式
    do_docx = args.docx
    do_pdf = args.pdf
    if not do_docx and not do_pdf:
        do_docx = True
        do_pdf = True

    base, _ = os.path.splitext(md_path)
    if args.output:
        out_path = os.path.abspath(args.output)
        out_base, out_ext = os.path.splitext(out_path)
        pdf_path = out_base + ".pdf"
        docx_path = out_base + ".docx"
    else:
        pdf_path = base + ".pdf"
        docx_path = base + ".docx"

    # 1. 生成 DOCX
    if do_docx:
        try:
            print(f"Generating ZUEL-styled DOCX at: {docx_path}...")
            generate_docx_from_markdown(md_path, docx_path)
            print(f"Success! ZUEL DOCX generated successfully at: {docx_path}")
        except Exception as e:
            print(f"Error during DOCX conversion: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            if not do_pdf:  # 如果只生成 docx 且失败了，退出
                sys.exit(1)

    # 2. 生成 PDF
    if do_pdf:
        if not os.path.exists(EDGE_PATH):
            print(f"Error: Microsoft Edge not found at: {EDGE_PATH}", file=sys.stderr)
            print("Please ensure Microsoft Edge is installed or adjust EDGE_PATH in the script.", file=sys.stderr)
            sys.exit(1)

        html_path = base + "_temp.html"
        try:
            html_data = generate_html(
                md_path=md_path,
                margin_top=args.margin_top,
                margin_bottom=args.margin_bottom,
                margin_left=args.margin_left,
                margin_right=args.margin_right
            )
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_data)

            cmd = [
                EDGE_PATH,
                "--headless=new",
                "--disable-gpu",
                f"--print-to-pdf={pdf_path}",
                "--no-pdf-header-footer",
                "--print-to-pdf-no-header",
                "--no-margins",
                html_path
            ]
            print("Rendering HTML and invoking Edge to print PDF...")
            subprocess.run(cmd, capture_output=True, text=True, check=True)
            print(f"Success! ZUEL PDF generated successfully at: {pdf_path}")

            # 后处理添加页脚页码和页眉
            try:
                post_process_pdf(
                    pdf_path=pdf_path,
                    header_text=args.header,
                    margin_left=args.margin_left,
                    margin_right=args.margin_right
                )
            except Exception as e:
                print(f"Warning: Failed to add page numbers/headers to PDF: {e}", file=sys.stderr)

        except Exception as e:
            print(f"Error during PDF conversion: {e}", file=sys.stderr)
            sys.exit(1)
        finally:
            if os.path.exists(html_path):
                os.remove(html_path)


def preview_cmd(args):
    """preview 子命令处理函数"""
    md_path = os.path.abspath(args.input)
    if not os.path.exists(md_path):
        print(f"Error: Input Markdown file does not exist: {md_path}", file=sys.stderr)
        sys.exit(1)
        
    if args.output:
        html_path = os.path.abspath(args.output)
    else:
        base, _ = os.path.splitext(md_path)
        html_path = base + "_preview.html"
        
    try:
        html_data = generate_html(
            md_path=md_path,
            margin_top=args.margin_top,
            margin_bottom=args.margin_bottom,
            margin_left=args.margin_left,
            margin_right=args.margin_right
        )
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_data)
        print(f"Success! Styled preview HTML generated at: {html_path}")
        
    except Exception as e:
        print(f"Error during preview generation: {e}", file=sys.stderr)
        sys.exit(1)

def validate_docx_internal(file_path):
    """内部校验 docx 并返回结构化数据与报告"""
    import docx
    import re
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
    doc = docx.Document(file_path)
    issues = []
    successes = []
    
    # 1. 校验页边距 (ZUEL: top=3.0cm, bottom=2.5cm, left=3.0cm, right=2.5cm)
    for idx, section in enumerate(doc.sections):
        top_cm = round(section.top_margin.cm, 2) if section.top_margin else None
        bottom_cm = round(section.bottom_margin.cm, 2) if section.bottom_margin else None
        left_cm = round(section.left_margin.cm, 2) if section.left_margin else None
        right_cm = round(section.right_margin.cm, 2) if section.right_margin else None
        
        if top_cm and abs(top_cm - 3.0) > 0.1:
            issues.append(f"第 {idx} 节上边距为 {top_cm}cm，标准为 3.0cm。")
        else:
            successes.append(f"第 {idx} 节上边距合规 ({top_cm}cm)。")
            
        if bottom_cm and abs(bottom_cm - 2.5) > 0.1:
            issues.append(f"第 {idx} 节下边距为 {bottom_cm}cm，标准为 2.5cm。")
        else:
            successes.append(f"第 {idx} 节下边距合规 ({bottom_cm}cm)。")
            
        if left_cm and abs(left_cm - 3.0) > 0.1:
            issues.append(f"第 {idx} 节左边距为 {left_cm}cm，标准为 3.0cm。")
        else:
            successes.append(f"第 {idx} 节左边距合规 ({left_cm}cm)。")
            
        if right_cm and abs(right_cm - 2.5) > 0.1:
            issues.append(f"第 {idx} 节右边距为 {right_cm}cm，标准为 2.5cm。")
        else:
            successes.append(f"第 {idx} 节右边距合规 ({right_cm}cm)。")

    # 获取字符字号的辅助函数
    def get_run_details(run):
        size = run.font.size.pt if run.font.size else None
        if size is None and run.style and run.style.font and run.style.font.size:
            size = run.style.font.size.pt
        font_name = run.font.name
        if font_name is None and run.style and run.style.font:
            font_name = run.style.font.name
        return size, font_name

    # 封面/章节合规标记
    has_logo_image = False
    has_school_name = False
    has_report_title = False
    has_course_info = 0
    has_abstract_title = False
    has_abstract_en_title = False
    has_keywords = False
    has_keywords_en = False
    has_toc = False
    has_references = False

    # 1.5 扫描所有表格中的课程信息
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p_cell in cell.paragraphs:
                    cell_text = p_cell.text.strip()
                    if not cell_text:
                        continue
                    for keyword in ["课程编号", "课程名称", "授课教师", "学生学号", "学生姓名"]:
                        if keyword in cell_text:
                            has_course_info += 1
                            runs_cell = p_cell.runs
                            p_sizes_cell = [r.font.size.pt for r in runs_cell if r.font.size]
                            if not p_sizes_cell and p_cell.style and p_cell.style.font and p_cell.style.font.size:
                                p_sizes_cell = [p_cell.style.font.size.pt]
                            avg_size_cell = sum(p_sizes_cell)/len(p_sizes_cell) if p_sizes_cell else None
                            if avg_size_cell and abs(avg_size_cell - 20) > 2:
                                issues.append(f"封面表格信息行 '{cell_text}' 字号为 {avg_size_cell}pt，标准为 20pt (二号/小二)。")
                            else:
                                successes.append(f"找到合规的封面表格信息行: '{cell_text}'。")
    
    in_cover = True
    in_toc_section = False
    seen_first_level_headings = set()

    # 2. 扫描所有段落
    for idx, p in enumerate(doc.paragraphs):
        # A. 检查封面校徽图片
        if in_cover and ('w:drawing' in p._p.xml or 'w:pict' in p._p.xml):
            has_logo_image = True
            
        text = p.text.strip()
        if not text:
            continue
            
        # B. 检查“中南财经政法大学” (仅在封面区域检测，精确匹配)
        if in_cover and text == "中南财经政法大学":
            has_school_name = True
            runs = p.runs
            p_bold = any(run.bold for run in runs)
            p_sizes = []
            for r in runs:
                sz, _ = get_run_details(r)
                if sz: p_sizes.append(sz)
            avg_size = sum(p_sizes)/len(p_sizes) if p_sizes else None
            if avg_size and abs(avg_size - 48) > 2:
                issues.append(f"校名 '{text}' 字号为 {avg_size}pt，标准为 48pt (初号)。")
            elif not p_bold:
                issues.append(f"校名 '{text}' 未加粗，标准应加粗。")
            else:
                successes.append(f"找到合规的校名: '{text}'。")
                
        # C. 检查“本科生结课报告” (仅在封面区域检测，精确匹配)
        if in_cover and text == "本科生结课报告":
            has_report_title = True
            runs = p.runs
            p_bold = any(run.bold for run in runs)
            p_sizes = []
            for r in runs:
                sz, _ = get_run_details(r)
                if sz: p_sizes.append(sz)
            avg_size = sum(p_sizes)/len(p_sizes) if p_sizes else None
            if avg_size and abs(avg_size - 48) > 2:
                issues.append(f"报告标题 '{text}' 字号为 {avg_size}pt，标准为 48pt (初号)。")
            elif not p_bold:
                issues.append(f"报告标题 '{text}' 未加粗，标准应加粗。")
            else:
                successes.append(f"找到合规的报告标题: '{text}'。")
            
        clean_text = "".join(text.split())
        is_目录_title = (clean_text == "目录")
        
        # 摘要页为前置页的起点，一旦遇到则判定封面结束
        if clean_text in ["摘要", "Abstract"]:
            in_cover = False
            
        runs = p.runs
        p_bold = any(run.bold for run in runs)
        p_sizes = []
        p_fonts = []
        for r in runs:
            sz, fn = get_run_details(r)
            if sz: p_sizes.append(sz)
            if fn: p_fonts.append(fn)
            
        avg_size = sum(p_sizes)/len(p_sizes) if p_sizes else None
        
        # 如果已经在目录区域中，且遇到了第一个一级标题，判断是否为正文开始
        if in_toc_section and not is_目录_title:
            if re.match(r"^[一二三四五六七八九十]+、", text):
                is_bold_heading = p_bold or (avg_size and avg_size >= 14) or p.paragraph_format.page_break_before
                if text in seen_first_level_headings or is_bold_heading:
                    in_toc_section = False
                else:
                    seen_first_level_headings.add(text)

        # 识别并跳过目录项，避免误判为标题
        is_toc_entry = (in_toc_section and not is_目录_title) or "TOC" in p.style.name or p.style.name.startswith("toc") or p.style.name.startswith("TOC")
        if not is_toc_entry:
            if (p.style.name == "Normal" or "Hyperlink" in p.style.name or "List" in p.style.name) and (re.search(r"\d+$", text) or "..." in text or "\t" in text or "•" in text):
                is_toc_entry = True
                
        if is_toc_entry:
            continue
            
        # 在检查完 is_toc_entry 之后，如果当前是目录标题，则设置 in_toc_section = True
        if is_目录_title:
            in_toc_section = True
        

                
        # D. 检查段落中的课程信息 (仅在封面区域检测)
        if in_cover:
            for keyword in ["课程编号", "课程名称", "授课教师", "学生学号", "学生姓名"]:
                if keyword in text:
                    has_course_info += 1
                    if avg_size and abs(avg_size - 20) > 2:
                        issues.append(f"封面信息行 '{text}' 字号为 {avg_size}pt，标准为 20pt (二号/小二)。")
                    else:
                        successes.append(f"找到合规的封面信息行: '{text}'。")

        # E. 检查中文/英文/结语/摘要/目录等大标题与关键词
        clean_text = "".join(text.split())
        if clean_text == "摘要":
            has_abstract_title = True
            if avg_size and abs(avg_size - 16) > 2:
                issues.append(f"摘要标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"摘要标题 '{text}' 未加粗，标准应加粗。")
            else:
                successes.append("中文摘要标题合规。")
                
        elif clean_text == "Abstract":
            has_abstract_en_title = True
            if avg_size and abs(avg_size - 16) > 2:
                issues.append(f"英文摘要标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"英文摘要标题 '{text}' 未加粗，标准应加粗。")
            else:
                successes.append("英文摘要标题合规。")
                
        elif clean_text == "结语":
            if avg_size and abs(avg_size - 16) > 2:
                issues.append(f"结语标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"结语标题 '{text}' 未加粗。")
            else:
                successes.append("结语标题合规。")
                
        if "关键词" in text:
            has_keywords = True
            lbl_run = next((r for r in runs if "关键词" in r.text), None)
            if lbl_run:
                sz, _ = get_run_details(lbl_run)
                if sz and abs(sz - 16) > 2:
                    issues.append(f"关键词标签 '{lbl_run.text}' 字号为 {sz}pt，标准为 16pt (三号)。")
                elif not lbl_run.bold:
                    issues.append(f"关键词标签 '{lbl_run.text}' 未加粗。")
                else:
                    successes.append("关键词标签合规。")
                    
        if "Keywords" in text or "Key words" in text:
            has_keywords_en = True
            lbl_run = next((r for r in runs if "Key" in r.text or "word" in r.text), None)
            if lbl_run:
                sz, _ = get_run_details(lbl_run)
                if sz and abs(sz - 16) > 2:
                    issues.append(f"英文关键词标签 '{lbl_run.text}' 字号为 {sz}pt，标准为 16pt (三号)。")
                elif not lbl_run.bold:
                    issues.append(f"英文关键词标签 '{lbl_run.text}' 未加粗。")
                else:
                    successes.append("英文关键词标签合规。")

        # F. 检查目录标题
        if clean_text == "目录":
            has_toc = True
            if avg_size and abs(avg_size - 16) > 2:
                issues.append(f"目录标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"目录标题 '{text}' 未加粗。")
            else:
                successes.append("目录标题合规。")

        # G. 检查参考文献标题
        if "主要参考文献" in text:
            has_references = True
            if avg_size and abs(avg_size - 16) > 2:
                issues.append(f"参考文献标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"参考文献标题 '{text}' 未加粗。")
            else:
                successes.append("参考文献标题合规。")

        # H. 检查正文标题级次
        # 一级标题 "一、"
        if re.match(r"^[一二三四五六七八九十]+、", text):
            # 必须是 黑体, 16pt, 加粗
            is_hei = any("黑" in str(f) for f in p_fonts) or p.style.name == "Heading 1"
            if avg_size and abs(avg_size - 16) > 1:
                issues.append(f"一级标题 '{text}' 字号为 {avg_size}pt，标准为 16pt (三号)。")
            elif not p_bold:
                issues.append(f"一级标题 '{text}' 未加粗。")
            else:
                successes.append(f"一级标题 '{text}' 格式合规。")
                
        # 二级标题 "（一）"
        elif re.match(r"^（[一二三四五六七八九十]+）", text):
            # 必须是 14pt, 加粗 (非黑体)
            if avg_size and abs(avg_size - 14) > 1:
                issues.append(f"二级标题 '{text}' 字号为 {avg_size}pt，标准为 14pt (四号)。")
            elif not p_bold:
                issues.append(f"二级标题 '{text}' 未加粗。")
            else:
                successes.append(f"二级标题 '{text}' 格式合规。")
                
        # 三级标题 "1."
        elif re.match(r"^\d+[\.、]", text):
            # 必须是 12pt, 加粗
            if avg_size and abs(avg_size - 12) > 1:
                issues.append(f"三级标题 '{text}' 字号为 {avg_size}pt，标准为 12pt (小四号)。")
            elif not p_bold:
                issues.append(f"三级标题 '{text}' 未加粗。")
            else:
                successes.append(f"三级标题 '{text}' 格式合规。")
                
        # I. 检查正文段落缩进与大小
        elif (p.style.name == "Normal" or "正文" in p.style.name) and len(text) > 20 and not text.startswith("……") and "引言" not in text and "摘要" not in text and "目录" not in text and "参考文献" not in text and "关键词" not in text and "Keywords" not in text and "Key words" not in text and "Key Words" not in text and not text.startswith("表") and not text.startswith("图"):
            # 必须是 12pt (小四)
            if avg_size and abs(avg_size - 12) > 1:
                issues.append(f"正文段落 '{text[:15]}...' 字号为 {avg_size}pt，标准为 12pt (小四)。")
            
            # 首行缩进检查
            indent = p.paragraph_format.first_line_indent
            if indent is None and p.style and p.style.paragraph_format:
                indent = p.style.paragraph_format.first_line_indent
            starts_with_spaces = text.startswith("  ") or text.startswith("    ") or text.startswith("\u3000\u3000")
            if indent is None and not starts_with_spaces:
                issues.append(f"正文段落 '{text[:15]}...' 首行未缩进 2 字符。")
            else:
                successes.append(f"正文段落 '{text[:15]}...' 缩进检查通过。")

    # 1.6 扫描并验证表格上方的题注 (表1 XXX - 去点化、居中、五号)
    for table in doc.tables:
        tbl_element = table._tbl
        parent = tbl_element.getparent()
        children = list(parent)
        try:
            tbl_idx = children.index(tbl_element)
            caption_p = None
            for idx_p in range(tbl_idx - 1, -1, -1):
                prev_element = children[idx_p]
                if prev_element.tag.endswith('p'):
                    from docx.text.paragraph import Paragraph
                    caption_p = Paragraph(prev_element, doc)
                    break
            
            if caption_p:
                cap_text = caption_p.text.strip()
                if cap_text.startswith("表") and any(cap_text.startswith(f"表{i}") or cap_text.startswith(f"表 {i}") for i in range(10)):
                    m_valid = re.match(r'^表\d+ \S+', cap_text)
                    if not m_valid:
                        issues.append(f"表格题注 '{cap_text}' 格式不合规，标准格式为 '表1 XXX'（表与数字间无空格，数字与文字间有且仅有一个空格分隔，且无点号）。")
                    else:
                        align = caption_p.alignment
                        if align is None and caption_p.style and caption_p.style.paragraph_format:
                            align = caption_p.style.paragraph_format.alignment
                        is_centered = (align == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER or align == 1)
                        
                        runs_cap = caption_p.runs
                        p_sizes_cap = []
                        for r in runs_cap:
                            sz, _ = get_run_details(r)
                            if sz: p_sizes_cap.append(sz)
                        avg_sz_cap = sum(p_sizes_cap)/len(p_sizes_cap) if p_sizes_cap else None
                        
                        if not is_centered:
                            issues.append(f"表格题注 '{cap_text}' 未居中对齐。")
                        elif avg_sz_cap and abs(avg_sz_cap - 10.5) > 1.5:
                            issues.append(f"表格题注 '{cap_text}' 字号为 {avg_sz_cap}pt，标准为 10.5pt (五号)。")
                        else:
                            successes.append(f"表格题注 '{cap_text}' 格式与样式合规。")
        except Exception as e:
            pass

    # 1.7 扫描并验证图片下方的题注 (图1 XXX - 去点化、居中、五号)
    for idx_p, p in enumerate(doc.paragraphs):
        if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
            if idx_p + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[idx_p + 1]
                next_text = next_p.text.strip()
                if next_text.startswith("图") and any(next_text.startswith(f"图{i}") or next_text.startswith(f"图 {i}") for i in range(10)):
                    m_valid = re.match(r'^图\d+ \S+', next_text)
                    if not m_valid:
                        issues.append(f"图片题注 '{next_text}' 格式不合规，标准格式为 '图1 XXX'（图与数字间无空格，数字与文字间有且仅有一个空格分隔，且无点号）。")
                    else:
                        align = next_p.alignment
                        if align is None and next_p.style and next_p.style.paragraph_format:
                            align = next_p.style.paragraph_format.alignment
                        is_centered = (align == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER or align == 1)
                        
                        runs_cap = next_p.runs
                        p_sizes_cap = []
                        for r in runs_cap:
                            sz, _ = get_run_details(r)
                            if sz: p_sizes_cap.append(sz)
                        avg_sz_cap = sum(p_sizes_cap)/len(p_sizes_cap) if p_sizes_cap else None
                        
                        if not is_centered:
                            issues.append(f"图片题注 '{next_text}' 未居中对齐。")
                        elif avg_sz_cap and abs(avg_sz_cap - 10.5) > 1.5:
                            issues.append(f"图片题注 '{next_text}' 字号为 {avg_sz_cap}pt，标准为 10.5pt (五号)。")
                        else:
                            successes.append(f"图片题注 '{next_text}' 格式与样式合规。")

    # 1.8 查找第一个一级标题并对其前面的大标题进行校验 (排除目录项)
    first_h1_idx = None
    in_toc_section = False
    seen_first_level_headings = set()
    for idx_p, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        if not p_text:
            continue
        clean_p = "".join(p_text.split())
        if clean_p == "目录":
            in_toc_section = True
            continue
        if in_toc_section:
            if re.match(r"^[一二三四五六七八九十]+、", p_text):
                if p_text in seen_first_level_headings:
                    in_toc_section = False
                    first_h1_idx = idx_p
                    break
                else:
                    seen_first_level_headings.add(p_text)
        else:
            if re.match(r"^[一二三四五六七八九十]+、", p_text):
                first_h1_idx = idx_p
                break
            
    if first_h1_idx is not None and first_h1_idx > 0:
        title_p = doc.paragraphs[first_h1_idx - 1]
        title_text = title_p.text.strip()
        if title_text and not title_text.startswith("表") and not title_text.startswith("图"):
            t_runs = title_p.runs
            t_bold = any(run.bold for run in t_runs)
            t_sizes = [r.font.size.pt for r in t_runs if r.font.size]
            t_avg_size = sum(t_sizes)/len(t_sizes) if t_sizes else None
            t_align = title_p.alignment
            
            if t_avg_size and abs(t_avg_size - 16) > 2 and abs(t_avg_size - 18) > 2:
                issues.append(f"论文正文大标题 '{title_text[:15]}...' 字号为 {t_avg_size}pt，标准为 16pt (三号) 或 18pt (小二)。")
            elif not t_bold:
                issues.append(f"论文正文大标题 '{title_text[:15]}...' 未加粗。")
            elif t_align != docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER and t_align != 1:
                issues.append(f"论文正文大标题 '{title_text[:15]}...' 未居中对齐。")
            else:
                successes.append(f"论文正文大标题 '{title_text[:15]}...' 格式合规（居中、加粗、符合字号）。")

    # 结构完整性校验
    if not has_logo_image:
        issues.append("未在封面检测到校徽图片。")
    else:
        successes.append("在封面成功检测到校徽图片。")
    if not has_school_name:
        issues.append("未在封面检测到校名“中南财经政法大学”。")
    if not has_report_title:
        issues.append("未在封面检测到报告标题“本科生结课报告”。")
    if has_course_info < 5:
        issues.append(f"封面填空项信息不完整（课程编号/名称/教师/学号/姓名，当前仅检测到 {has_course_info}/5 项）。")
    if not has_abstract_title:
        issues.append("未检测到中文“摘要”标题。")
    if not has_abstract_en_title:
        issues.append("未检测到英文“Abstract”标题。")
    if not has_keywords:
        issues.append("未检测到“关键词：”标签。")
    if not has_keywords_en:
        issues.append("未检测到英文“Key words:”标签。")
    if not has_toc:
        issues.append("未检测到“目录”页。")
    if not has_references:
        issues.append("未检测到“主要参考文献”章节。")

    return {
        "file_path": file_path,
        "is_compliant": len(issues) == 0,
        "errors": issues,
        "successes": successes
    }

def validate_docx_cmd(args):
    """validate-docx 子命令入口"""
    docx_path = os.path.abspath(args.input)
    if not os.path.exists(docx_path):
        print(f"Error: Input DOCX file does not exist: {docx_path}", file=sys.stderr)
        sys.exit(1)
        
    print(f"Starting ZUEL formatting check for: {os.path.basename(docx_path)}...")
    try:
        result = validate_docx_internal(docx_path)
    except Exception as e:
        print(f"Error reading docx: {e}", file=sys.stderr)
        sys.exit(1)
        
    # 生成报告文件路径
    base, _ = os.path.splitext(docx_path)
    report_md_path = args.output_report if args.output_report else base + "_validation_report.md"
    report_json_path = args.output_json if args.output_json else base + "_validation_report.json"
    
    # 1. 写入 JSON 报告
    with open(report_json_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
        
    # 2. 构造 Markdown 报告
    status_str = "🟢 **完全合规 (Compliant)**" if result["is_compliant"] else "🔴 **存在缺陷 (Non-Compliant)**"
    errors_block = ""
    if result["errors"]:
        errors_block = "\n".join([f"- [ ] ❌ {err}" for err in result["errors"]])
    else:
        errors_block = "🎉 恭喜！未检测到任何格式不合规项，文档格式完美契合 ZUEL 模版要求。"
        
    checklist_rows = []
    for s in result["successes"]:
        checklist_rows.append(f"| {s} | 🟢 通过 |")
    for e in result["errors"]:
        checklist_rows.append(f"| {e} | 🔴 失败 |")
    checklist_table = "\n".join(checklist_rows)
    
    md_content = f"""# ZUEL 本科生结课论文格式校验报告

## 📊 校验概览
- **校验文件**: `{os.path.basename(docx_path)}`
- **校验结论**: {status_str}
- **合规项总数**: `{len(result["successes"])}` 项
- **缺陷项总数**: `{len(result["errors"])}` 项

## ❌ 待修正的缺陷清单
{errors_block}

## 🔍 详细检查清单
| 校验内容描述 | 校验结果 |
| :--- | :---: |
{checklist_table}

---
*报告由 ZUEL-Formatter 自动生成。*
"""
    
    with open(report_md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
        
    print("\n=== 校验结果汇总 ===")
    print(f"合规项数: {len(result['successes'])}")
    print(f"缺陷项数: {len(result['errors'])}")
    print(f"校验状态: {'PASS (合规)' if result['is_compliant'] else 'FAIL (不合规)'}")
    print(f"详细 Markdown 报告已写入: {report_md_path}")
    print(f"详细 JSON 数据报告已写入: {report_json_path}")
    
    if not result["is_compliant"]:
        print("\n❌ 发现以下不合规项:")
        for err in result["errors"][:5]:
            print(f"  - {err}")
        if len(result["errors"]) > 5:
            print(f"  ... 还有 {len(result['errors'])-5} 项，请打开报告查看完整清单。")
            
    # 如果不合规，退出代码为 0 (正常结束但报告结果)，或者也可以视具体要求设定
    sys.exit(0)

def main():
    parser = argparse.ArgumentParser(description="ZUEL-Formatter: Convert MD to A4 PDF, or validate student docx formatting.")
    subparsers = parser.add_subparsers(dest="command", help="Subcommands")
    subparsers.required = True
    
    # convert 命令
    convert_parser = subparsers.add_parser("convert", help="Convert MD to academic ZUEL A4 PDF and/or DOCX formats")
    convert_parser.add_argument("--input", "-i", required=True, help="Input Markdown file path")
    convert_parser.add_argument("--output", "-o", help="Output file path (optional, will auto-determine .pdf and .docx extensions)")
    convert_parser.add_argument("--docx", action="store_true", help="Generate ZUEL-styled DOCX only (or with PDF if --pdf is also specified)")
    convert_parser.add_argument("--pdf", action="store_true", help="Generate ZUEL-styled PDF only (or with DOCX if --docx is also specified)")
    convert_parser.add_argument("--header", help="Optional text for top page header")
    convert_parser.add_argument("--margin-top", type=float, default=3.0, help="Top margin in cm (default: 3.0)")
    convert_parser.add_argument("--margin-bottom", type=float, default=2.5, help="Bottom margin in cm (default: 2.5)")
    convert_parser.add_argument("--margin-left", type=float, default=3.0, help="Left margin in cm (default: 3.0)")
    convert_parser.add_argument("--margin-right", type=float, default=2.5, help="Right margin in cm (default: 2.5)")
    convert_parser.set_defaults(func=convert_cmd)
    
    # preview 命令
    preview_parser = subparsers.add_parser("preview", help="Generate styled HTML preview page")
    preview_parser.add_argument("--input", "-i", required=True, help="Input Markdown file path")
    preview_parser.add_argument("--output", "-o", help="Output HTML file path (optional)")
    preview_parser.add_argument("--header", help="Optional text for top page header")
    preview_parser.add_argument("--margin-top", type=float, default=3.0, help="Top margin in cm (default: 3.0)")
    preview_parser.add_argument("--margin-bottom", type=float, default=2.5, help="Bottom margin in cm (default: 2.5)")
    preview_parser.add_argument("--margin-left", type=float, default=3.0, help="Left margin in cm (default: 3.0)")
    preview_parser.add_argument("--margin-right", type=float, default=2.5, help="Right margin in cm (default: 2.5)")
    preview_parser.set_defaults(func=preview_cmd)
    
    # validate-docx 命令
    validate_parser = subparsers.add_parser("validate-docx", help="Validate DOCX file against ZUEL formatting criteria")
    validate_parser.add_argument("--input", "-i", required=True, help="Input DOCX file path to validate")
    validate_parser.add_argument("--output-report", "-o", help="Output Markdown report path (optional)")
    validate_parser.add_argument("--output-json", "-j", help="Output JSON report path (optional)")
    validate_parser.set_defaults(func=validate_docx_cmd)
    
    args = parser.parse_args()
    args.func(args)

if __name__ == '__main__':
    main()
